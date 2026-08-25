try:
    import pymupdf as fitz
except ImportError:
    import fitz

from ..utils.text_cleaner import clean_text, extract_sections


# ---------------------------------------------------------------------------
# Low-level extractors: bytes -> raw text
# ---------------------------------------------------------------------------

def _pdf_bytes_to_text(data: bytes) -> str:
    import tempfile
    import os

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        doc = fitz.open(tmp_path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
    finally:
        os.unlink(tmp_path)
    return text


def _docx_bytes_to_text(data: bytes) -> str:
    """Extract paragraphs and table cells from a .docx resume."""
    from docx import Document  # python-docx
    import io
    import tempfile
    import os

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    finally:
        os.unlink(tmp_path)


IMAGE_MIMES = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
    "bmp": "image/bmp",
    "gif": "image/gif",
}


def _image_bytes_to_text(data: bytes, ext: str) -> str:
    """Gemini vision OCR for photo/scan resumes."""
    from ..ai.gemini_client import extract_text_from_image

    mime = IMAGE_MIMES.get(ext.lower(), "image/jpeg")
    return extract_text_from_image(data, mime_type=mime)


# ---------------------------------------------------------------------------
# Text -> structured resume data
# ---------------------------------------------------------------------------

def _text_to_resume_data(raw_text, source_label="") -> dict:
    if not raw_text or not raw_text.strip():
        raise ValueError(
            f"Could not extract readable text from this {source_label or 'file'}. "
            "If it is a scanned document, upload a clear photo/image of it instead."
        )
    cleaned = clean_text(raw_text)
    sections = extract_sections(cleaned)
    return {"raw_text": raw_text, "clean_text": cleaned, "sections": sections}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_resume_pdf(file_path) -> dict:
    """Parse a resume PDF from a filesystem path."""
    doc = fitz.open(file_path)
    full_text = "\n".join(page.get_text() for page in doc)
    doc.close()
    return _text_to_resume_data(full_text, "PDF")


def parse_resume_txt(raw_bytes) -> dict:
    """Parse a plain-text resume."""
    try:
        text = raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = raw_bytes.decode("latin-1", errors="ignore")
    if not text.strip():
        raise ValueError("The uploaded text file is empty.")
    return _text_to_resume_data(text, "TXT")


def parse_resume_docx(raw_bytes) -> dict:
    """Parse a Word (.docx) resume."""
    return _text_to_resume_data(_docx_bytes_to_text(raw_bytes), "DOCX")


def parse_resume_image(raw_bytes, ext: str) -> dict:
    """Parse a resume image (png/jpg/webp...) via Gemini vision OCR.

    Raises a clear error when no API key is configured so the user knows why.
    """
    from ..ai.gemini_client import GEMINI_API_KEY

    if not GEMINI_API_KEY:
        raise ValueError(
            "Image parsing requires GEMINI_API_KEY. Add it to .env "
            "or upload a PDF / DOCX / TXT instead."
        )
    text = _image_bytes_to_text(raw_bytes, ext)
    return _text_to_resume_data(text, "image")


SUPPORTED_EXTENSIONS = ["pdf", "txt", "docx", "doc", "png", "jpg", "jpeg", "webp", "bmp"]


def parse_resume_upload(uploaded_file) -> dict:
    """Route an uploaded resume by extension to the right extractor."""
    name = (getattr(uploaded_file, "name", "") or "").lower()
    data = uploaded_file.read()

    ext = name.rsplit(".", 1)[-1] if "." in name else ""

    if ext == "pdf":
        return _from_bytes(data)
    if ext == "txt":
        return parse_resume_txt(data)
    if ext == "docx":
        return parse_resume_docx(data)
    if ext == "doc":
        raise ValueError(
            "Legacy .doc is not supported. Please save as .docx or export to PDF."
        )
    if ext in IMAGE_MIMES:
        return parse_resume_image(data, ext)

    # A PDF signature is unambiguous even when the filename is missing.
    if data[:5] == b"%PDF-":
        return _from_bytes(data)
    raise ValueError(
        "Unsupported resume format. Please upload a PDF, TXT, DOCX, PNG, JPG, or WEBP file."
    )


def _from_bytes(pdf_bytes: bytes) -> dict:
    return _text_to_resume_data(_pdf_bytes_to_text(pdf_bytes), "PDF")
